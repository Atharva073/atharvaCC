from django.shortcuts import render,redirect
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 
from django.contrib.auth.hashers import check_password
from .models import Login
from django.contrib.auth.hashers import make_password


# Create your views here.
#home
def home(request):
    return render(request, 'pc/registration.html') 

def index_view(request):
    return render(request, 'pc/index.html')

def registration_view(request):
    if request.method == 'POST':
        # Get form data
        username = request.POST.get('username')
        password = make_password(request.POST.get('password'))  
        name = request.POST.get('name')
        roll_number = request.POST.get('roll_number')
        login = Login.objects.create(username=username, password=password, name=name, roll_number=roll_number)
        return HttpResponse('Registration successful!')  # For demonstration purposes
    return render(request, 'pc/registration.html')

def loginView(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        try:
            user = Login.objects.get(username=username)
            if check_password(password, user.password):
                return render(request,'pc/index.html')  # Redirect to the home page after successful login
            else:
                # Invalid password
                return HttpResponse('Invalid password!')  # For demonstration purposes
                #return render(request, 'pc/login.html', {'error_message': 'Invalid username or password'})
        except Login.DoesNotExist:
            return HttpResponse('User not found !')  # For demonstration purposes

    return render(request, 'pc/login.html')

#web search(Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    print(request.POST['q1'])  
    print(request.POST['q2'])  
    
    p1=0
    p2=0
    p3=0
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        p1 = round(percent,2)
    if request.POST['q1']: 
        percent,link = main.findSimilarity(request.POST['q1'])
        p2 = round(percent,2)
    if request.POST['q2']: 
        percent,link = main.findSimilarity(request.POST['q2'])
        p3 = round(percent,2)
    print("Output.....................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent,'p1':p1,'p2':p2,'p3':p3,})

#web search file(.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object 
        pdfFileObj = open(request.FILES['docfile'], 'rb') 

        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # printing number of pages in pdf file 
        print(pdfReader.numPages) 

        # creating a page object 
        pageObj = pdfReader.getPage(0) 

        # extracting text from page 
        print(pageObj.extractText()) 

        # closing the pdf file object 
        pdfFileObj.close() 


    percent,link = main.findSimilarity(value)
    print("Output...................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    
    print("Output..................!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
